"""DOCX document renderer."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from lewlm.core.errors import DocumentGenerationError
from lewlm.documents.ir.models import (
    CalloutBlock,
    DocumentIR,
    DocumentOutputFormat,
    ImageBlock,
    ListBlock,
    ParagraphBlock,
    TableBlock,
)
from lewlm.documents.render.base import DocumentRenderer
from lewlm.documents.validators.ir import DocumentIRValidator


class DocxDocumentRenderer(DocumentRenderer):
    output_format = DocumentOutputFormat.DOCX
    media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    file_extension = ".docx"

    def __init__(self, validator: DocumentIRValidator) -> None:
        self.validator = validator

    def render(self, document: DocumentIR) -> bytes:
        self.validator.validate(document)
        try:
            from docx import Document as WordDocument
        except ImportError as exc:
            raise DocumentGenerationError(
                "DOCX generation requires the `python-docx` dependency.",
                details={"renderer": self.output_format.value},
            ) from exc

        word_document = WordDocument()
        word_document.core_properties.title = document.title
        self._render_header_footer(word_document, document)
        word_document.add_heading(document.title, level=0)

        for section in document.sections:
            if section.heading:
                word_document.add_heading(section.heading, level=max(1, min(section.level, 9)))
            for block in section.blocks:
                if isinstance(block, ParagraphBlock):
                    word_document.add_paragraph(block.text)
                elif isinstance(block, ListBlock):
                    style = "List Number" if block.ordered else "List Bullet"
                    for item in block.items:
                        word_document.add_paragraph(item, style=style)
                elif isinstance(block, TableBlock):
                    self._render_table(word_document, block)
                elif isinstance(block, CalloutBlock):
                    paragraph = word_document.add_paragraph()
                    prefix = f"{block.kind.upper()}"
                    if block.title:
                        prefix = f"{prefix}: {block.title}"
                    paragraph.add_run(prefix).bold = True
                    paragraph.add_run(f" {block.body}")
                elif isinstance(block, ImageBlock):
                    self._render_image(word_document, block)

        if document.citations:
            word_document.add_heading(document.references_title, level=1)
            for citation in document.citations:
                text = citation.text if citation.url is None else f"{citation.text} ({citation.url})"
                word_document.add_paragraph(f"[{citation.label}] {text}")

        buffer = BytesIO()
        word_document.save(buffer)
        return buffer.getvalue()

    def _render_header_footer(self, word_document, document: DocumentIR) -> None:
        section = word_document.sections[0]
        if document.header:
            header = section.header.paragraphs[0]
            header.text = " | ".join(value for value in (document.header.left, document.header.center, document.header.right) if value)
        if document.footer:
            footer = section.footer.paragraphs[0]
            footer.text = " | ".join(value for value in (document.footer.left, document.footer.center, document.footer.right) if value)

    def _render_table(self, word_document, block: TableBlock) -> None:
        column_count = len(block.headers) if block.headers else len(block.rows[0])
        table = word_document.add_table(rows=0, cols=column_count)
        table.style = "Table Grid"
        if block.headers:
            header_cells = table.add_row().cells
            for index, header in enumerate(block.headers):
                header_cells[index].text = header
        for row in block.rows:
            row_cells = table.add_row().cells
            for index, value in enumerate(row):
                row_cells[index].text = value
        if block.caption:
            word_document.add_paragraph(block.caption)

    def _render_image(self, word_document, block: ImageBlock) -> None:
        if block.path is not None:
            word_document.add_picture(str(Path(block.path).expanduser()))
        else:
            word_document.add_paragraph(f"[Image] {block.alt_text}")
        if block.caption:
            word_document.add_paragraph(block.caption)
