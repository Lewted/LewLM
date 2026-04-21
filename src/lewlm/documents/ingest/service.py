"""Deterministic local document ingestion service."""

from __future__ import annotations

import csv
from dataclasses import dataclass
import hashlib
from io import BytesIO, StringIO
import mimetypes
from pathlib import Path
import re
import shutil
import zipfile
from typing import Any
from uuid import uuid4

from lewlm.core.errors import DocumentValidationError, PackUnavailableError, UnsupportedMediaTypeError
from lewlm.documents.ingest.models import DocumentChunk, DocumentIngestResult, DocumentSourceType, IngestedDocumentSource
from lewlm.documents.ingest.ocr import OcrBackendStatus, detect_ocr_backend, perform_ocr_on_image_bytes
from lewlm.documents.ir.models import CalloutBlock, DocumentIR, DocumentSection, ImageBlock, ListBlock, ParagraphBlock, TableBlock
from lewlm.events.bus import EventBus
from lewlm.events.schema import EventScope, EventType, StreamEvent
from lewlm.security.files import (
    read_scoped_text_file,
    resolve_scoped_path,
    validate_scoped_binary_file,
    validate_scoped_image_file,
)
from lewlm.security.sandbox import run_in_subprocess
from lewlm.security.workspace import secure_workspace


@dataclass(slots=True)
class _ParsedSource:
    sections: list[DocumentSection]
    source: IngestedDocumentSource


def _split_headers_and_rows(rows: list[list[str]]) -> tuple[list[str], list[list[str]]]:
    if len(rows) == 1:
        return [], rows
    return rows[0], rows[1:]


def _default_ingest_scope_roots(
    paths: list[str] | tuple[str, ...] | list[Path] | tuple[Path, ...],
    *,
    base_dir: Path | str | None,
) -> tuple[Path, ...]:
    roots: list[Path] = []
    for raw_path in paths:
        candidate = Path(raw_path).expanduser()
        resolved = (
            (Path(base_dir).expanduser().resolve(strict=False) / candidate).resolve(strict=False)
            if base_dir is not None and not candidate.is_absolute()
            else candidate.resolve(strict=False)
        )
        roots.append(resolved if resolved.is_dir() else resolved.parent)
    deduped_roots: list[Path] = []
    for root in roots:
        if root not in deduped_roots:
            deduped_roots.append(root)
    return tuple(deduped_roots)


def _parse_csv_file(parser_path: str) -> dict[str, Any]:
    rows = [
        [str(cell) for cell in row]
        for row in csv.reader(StringIO(Path(parser_path).read_text(encoding="utf-8")))
        if any(cell.strip() for cell in row)
    ]
    if not rows:
        raise DocumentValidationError("CSV document did not contain any populated rows.", details={"path": parser_path})
    headers, data_rows = _split_headers_and_rows(rows)
    return {
        "headers": headers,
        "rows": data_rows,
        "row_count": len(data_rows),
        "column_count": len(data_rows[0]) if data_rows else len(headers),
    }


def _parse_xlsx_file(parser_path: str) -> dict[str, Any]:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise DocumentValidationError("XLSX ingest requires the `openpyxl` dependency.") from exc

    workbook = load_workbook(parser_path, read_only=True, data_only=True)
    try:
        sections = []
        for sheet in workbook.worksheets:
            rows = []
            for row in sheet.iter_rows(values_only=True):
                values = ["" if value is None else str(value) for value in row]
                if any(value != "" for value in values):
                    rows.append(values)
            if not rows:
                continue
            headers, data_rows = _split_headers_and_rows(rows)
            sections.append({"title": sheet.title, "headers": headers, "rows": data_rows})
    finally:
        workbook.close()
    if not sections:
        raise DocumentValidationError("XLSX document did not contain any populated sheets.", details={"path": parser_path})
    return {"sheets": sections, "sheet_count": len(sections)}


def _parse_docx_file(parser_path: str) -> dict[str, Any]:
    try:
        from docx import Document as WordDocument
    except ImportError as exc:
        raise DocumentValidationError("DOCX ingest requires the `python-docx` dependency.") from exc

    word_document = WordDocument(parser_path)
    paragraphs = [
        {
            "text": paragraph.text.strip(),
            "style": getattr(paragraph.style, "name", None),
        }
        for paragraph in word_document.paragraphs
        if paragraph.text.strip()
    ]
    tables = []
    for table in word_document.tables:
        rows = [
            [cell.text.strip() for cell in row.cells]
            for row in table.rows
            if any(cell.text.strip() for cell in row.cells)
        ]
        if not rows:
            continue
        headers, data_rows = _split_headers_and_rows(rows)
        tables.append({"headers": headers, "rows": data_rows})
    return {
        "paragraphs": paragraphs,
        "tables": tables,
        "inline_image_count": len(getattr(word_document, "inline_shapes", [])),
    }


def _parse_pdf_file(parser_path: str) -> dict[str, Any]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise DocumentValidationError("PDF ingest requires the `pypdf` dependency.") from exc

    reader = PdfReader(parser_path)
    pages = []
    text_page_count = 0
    total_image_count = 0
    for page_index, page in enumerate(reader.pages, start=1):
        extracted_text = (page.extract_text() or "").strip()
        page_images = []
        for image_index, image in enumerate(page.images, start=1):
            image_name = getattr(image, "name", None) or f"page-{page_index}-image-{image_index}.bin"
            page_images.append({"name": image_name, "data": bytes(image.data)})
        total_image_count += len(page_images)
        if extracted_text:
            text_page_count += 1
            paragraphs = [segment.strip() for segment in re.split(r"\n\s*\n", extracted_text) if segment.strip()]
            pages.append({"paragraphs": paragraphs or [extracted_text], "images": page_images})
        else:
            pages.append({"paragraphs": [], "images": page_images, "warning": "No extractable text found on this page."})
    return {
        "pages": pages,
        "page_count": len(reader.pages),
        "text_page_count": text_page_count,
        "image_count": total_image_count,
    }


def _parse_text_file(parser_path: str) -> dict[str, Any]:
    paragraphs = _paragraphs_from_text(Path(parser_path).read_text(encoding="utf-8"))
    if not paragraphs:
        raise DocumentValidationError("TXT document did not contain any extractable text.", details={"path": parser_path})
    return {"paragraphs": paragraphs}


def _parse_markdown_file(parser_path: str) -> dict[str, Any]:
    text = Path(parser_path).read_text(encoding="utf-8")
    default_heading = Path(parser_path).stem
    sections: list[dict[str, Any]] = []
    current_heading = default_heading
    current_level = 1
    current_blocks: list[dict[str, Any]] = []
    paragraph_lines: list[str] = []
    list_items: list[str] = []
    list_ordered = False
    in_code_block = False
    code_lines: list[str] = []

    def flush_paragraph() -> None:
        if not paragraph_lines:
            return
        paragraph = _join_wrapped_lines(paragraph_lines)
        paragraph_lines.clear()
        if paragraph:
            current_blocks.append({"type": "paragraph", "text": paragraph})

    def flush_list() -> None:
        nonlocal list_ordered
        if not list_items:
            return
        current_blocks.append(
            {
                "type": "list",
                "ordered": list_ordered,
                "items": list(list_items),
            },
        )
        list_items.clear()
        list_ordered = False

    def flush_code_block() -> None:
        if not code_lines:
            return
        current_blocks.append(
            {
                "type": "paragraph",
                "text": "\n".join(code_lines).strip(),
                "style_tokens": ["code"],
            },
        )
        code_lines.clear()

    def flush_section() -> None:
        flush_paragraph()
        flush_list()
        if current_blocks:
            sections.append(
                {
                    "heading": current_heading,
                    "level": current_level,
                    "blocks": list(current_blocks),
                },
            )
            current_blocks.clear()

    for raw_line in text.splitlines():
        stripped = raw_line.strip()
        if stripped.startswith("```"):
            flush_paragraph()
            flush_list()
            if in_code_block:
                flush_code_block()
                in_code_block = False
            else:
                in_code_block = True
            continue
        if in_code_block:
            code_lines.append(raw_line.rstrip())
            continue
        if not stripped:
            flush_paragraph()
            flush_list()
            continue
        heading_match = re.match(r"^(#{1,6})\s+(.+?)(?:\s+#+\s*)?$", stripped)
        if heading_match is not None:
            flush_section()
            current_heading = heading_match.group(2).strip()
            current_level = len(heading_match.group(1))
            continue
        list_match = re.match(r"^(?P<marker>(?:[-*+])|(?:\d+\.))\s+(?P<item>.+)$", stripped)
        if list_match is not None:
            flush_paragraph()
            ordered = list_match.group("marker").endswith(".")
            if list_items and list_ordered != ordered:
                flush_list()
            list_ordered = ordered
            list_items.append(list_match.group("item").strip())
            continue
        flush_list()
        paragraph_lines.append(stripped)

    if in_code_block:
        flush_code_block()
    flush_section()
    if not sections:
        raise DocumentValidationError("Markdown document did not contain any extractable text.", details={"path": parser_path})

    paragraph_count = sum(
        1
        for section in sections
        for block in section["blocks"]
        if block["type"] == "paragraph"
    )
    list_count = sum(
        1
        for section in sections
        for block in section["blocks"]
        if block["type"] == "list"
    )
    return {
        "sections": sections,
        "section_count": len(sections),
        "paragraph_count": paragraph_count,
        "list_count": list_count,
    }


def _read_image_metadata(parser_path: str) -> dict[str, Any]:
    try:
        from PIL import Image, ImageOps
    except ImportError as exc:
        raise DocumentValidationError("Image ingest requires the `Pillow` dependency.") from exc

    with Image.open(parser_path) as image:
        normalized = ImageOps.exif_transpose(image)
        width, height = normalized.size
        image_format = image.format or Path(parser_path).suffix.lstrip(".").upper() or "UNKNOWN"
    return {"format": image_format, "width": width, "height": height}


def _read_image_metadata_bytes(image_bytes: bytes, file_name: str) -> dict[str, Any]:
    try:
        from PIL import Image, ImageOps
    except ImportError as exc:
        raise DocumentValidationError("Image ingest requires the `Pillow` dependency.") from exc

    with Image.open(BytesIO(image_bytes)) as image:
        normalized = ImageOps.exif_transpose(image)
        width, height = normalized.size
        image_format = image.format or Path(file_name).suffix.lstrip(".").upper() or "UNKNOWN"
    return {
        "file_name": Path(file_name).name,
        "format": image_format,
        "width": width,
        "height": height,
        "mime_type": _guess_image_media_type(file_name, image_format),
    }


class DocumentIngestService:
    """Safely ingest local documents into LewLM's structured IR."""

    def __init__(
        self,
        *,
        workspace_root: Path,
        sandbox_enabled: bool = True,
        sandbox_timeout_seconds: int = 30,
        sandbox_clear_environment: bool = True,
        chunk_max_characters: int = 1200,
        event_bus: EventBus | None = None,
        enabled: bool = True,
        disabled_reason: str | None = None,
    ) -> None:
        self.workspace_root = workspace_root
        self.sandbox_enabled = sandbox_enabled
        self.sandbox_timeout_seconds = sandbox_timeout_seconds
        self.sandbox_clear_environment = sandbox_clear_environment
        self.chunk_max_characters = chunk_max_characters
        self.event_bus = event_bus
        self.enabled = enabled
        self.disabled_reason = disabled_reason or "Document ingest is disabled for this LewLM process."

    def ingest(
        self,
        paths: list[str] | tuple[str, ...] | list[Path] | tuple[Path, ...],
        *,
        title: str | None = None,
        allowed_file_roots: tuple[Path | str, ...] | list[Path | str] | None = None,
        base_dir: Path | str | None = None,
        request_id: str | None = None,
    ) -> DocumentIngestResult:
        self._ensure_enabled()
        if not paths:
            raise DocumentValidationError("Document ingest requires at least one source path.")
        resolved_request_id = request_id or str(uuid4())
        scoped_roots = tuple(allowed_file_roots) if allowed_file_roots is not None else _default_ingest_scope_roots(paths, base_dir=base_dir)
        resolved_sources = [
            resolve_scoped_path(
                path,
                allowed_roots=scoped_roots,
                purpose="Document source",
                base_dir=base_dir,
                expect="any",
            )
            for path in paths
        ]
        ocr_status = detect_ocr_backend()
        self._publish(
            EventType.DOCUMENT_PARSE_STARTED,
            {
                "request_id": resolved_request_id,
                "source_count": len(resolved_sources),
                "title": title,
                "ocr_available": ocr_status.available,
            },
        )

        sections: list[DocumentSection] = []
        source_records: list[IngestedDocumentSource] = []
        try:
            with secure_workspace(self.workspace_root, prefix="ingest-") as workspace:
                total_sources = len(resolved_sources)
                for index, source_path in enumerate(resolved_sources):
                    parsed = self._ingest_source(
                        source_path,
                        workspace / f"source-{index}",
                        allowed_file_roots=scoped_roots,
                        base_dir=base_dir,
                        ocr_status=ocr_status,
                    )
                    source_record = _standardize_source_record(parsed.source, source_index=index)
                    parsed_sections = _standardize_sections(
                        parsed.sections,
                        source=source_record,
                        source_index=index,
                        global_section_offset=len(sections),
                    )
                    sections.extend(parsed_sections)
                    source_records.append(source_record)
                    self._publish(
                        EventType.OPERATION_PROGRESS,
                        {
                            "request_id": resolved_request_id,
                            "operation": "document.parse",
                            "stage": "source_parsed",
                            "completed_steps": index + 1,
                            "total_steps": total_sources,
                            "progress": round((index + 1) / total_sources, 4),
                            "path": str(source_path),
                            "source_type": source_record.source_type.value,
                        },
                    )

            document_title = title or self._default_title(resolved_sources)
            document = DocumentIR(
                title=document_title,
                metadata={
                    "source_count": len(source_records),
                    "source_ids": [source.source_id for source in source_records],
                    "source_labels": [source.source_label for source in source_records],
                    "source_types": [source.source_type.value for source in source_records],
                    "ocr_available": ocr_status.available,
                    "ocr_backend": ocr_status.backend_name,
                    "ocr_reason": ocr_status.reason,
                    "ocr_used": any(bool(source.metadata.get("ocr_used")) for source in source_records),
                    "chunk_strategy": "section_block_semantic_segmentation",
                },
                sections=sections,
            )
            chunks = self._build_chunks(document)
            document = document.model_copy(
                update={
                    "metadata": {
                        **document.metadata,
                        "chunk_count": len(chunks),
                    },
                },
            )
            self._publish(
                EventType.DOCUMENT_PARSE_COMPLETED,
                {
                    "request_id": resolved_request_id,
                    "title": document_title,
                    "source_count": len(source_records),
                    "section_count": len(sections),
                    "chunk_count": len(chunks),
                    "ocr_used": document.metadata["ocr_used"],
                },
            )
            return DocumentIngestResult(document=document, sources=source_records, chunks=chunks)
        except Exception as exc:
            self._publish(
                EventType.DOCUMENT_PARSE_FAILED,
                {
                    "request_id": resolved_request_id,
                    "source_count": len(resolved_sources),
                    "error": str(exc),
                },
            )
            raise

    def _ensure_enabled(self) -> None:
        if self.enabled:
            return
        raise PackUnavailableError(
            "Document ingest is unavailable because the `documents` feature pack is disabled.",
            details={"pack": "documents", "reason": self.disabled_reason},
        )

    def _ingest_source(
        self,
        source_path: Path,
        workspace_path: Path,
        *,
        allowed_file_roots: tuple[Path | str, ...],
        base_dir: Path | str | None,
        ocr_status: OcrBackendStatus,
    ) -> _ParsedSource:
        if source_path.is_dir():
            return self._ingest_image_bundle(
                source_path,
                workspace_path,
                allowed_file_roots=allowed_file_roots,
                base_dir=base_dir,
                ocr_status=ocr_status,
            )

        suffix = source_path.suffix.casefold()
        if suffix == ".txt":
            return self._ingest_text(
                source_path,
                workspace_path,
                allowed_file_roots=allowed_file_roots,
                base_dir=base_dir,
            )
        if suffix in {".md", ".markdown"}:
            return self._ingest_markdown(
                source_path,
                workspace_path,
                allowed_file_roots=allowed_file_roots,
                base_dir=base_dir,
            )
        if suffix == ".csv":
            return self._ingest_csv(source_path, workspace_path, allowed_file_roots=allowed_file_roots, base_dir=base_dir)
        if suffix == ".xlsx":
            return self._ingest_xlsx(source_path, workspace_path, allowed_file_roots=allowed_file_roots, base_dir=base_dir)
        if suffix == ".docx":
            return self._ingest_docx(
                source_path,
                workspace_path,
                allowed_file_roots=allowed_file_roots,
                base_dir=base_dir,
                ocr_status=ocr_status,
            )
        if suffix == ".pdf":
            return self._ingest_pdf(
                source_path,
                workspace_path,
                allowed_file_roots=allowed_file_roots,
                base_dir=base_dir,
                ocr_status=ocr_status,
            )
        return self._ingest_image(
            source_path,
            workspace_path,
            allowed_file_roots=allowed_file_roots,
            base_dir=base_dir,
            ocr_status=ocr_status,
        )

    def _ingest_csv(
        self,
        source_path: Path,
        workspace_path: Path,
        *,
        allowed_file_roots: tuple[Path | str, ...],
        base_dir: Path | str | None,
    ) -> _ParsedSource:
        resolved_path, _ = read_scoped_text_file(
            source_path,
            allowed_roots=allowed_file_roots,
            purpose="CSV document",
            media_type="text/plain",
            base_dir=base_dir,
        )
        parser_path = self._copy_to_workspace(resolved_path, workspace_path.with_suffix(".csv"))
        parsed = self._run_parser("CSV ingest parser", _parse_csv_file, str(parser_path))
        return _ParsedSource(
            sections=[
                DocumentSection(
                    heading=resolved_path.stem,
                    metadata={"source_path": str(resolved_path), "source_type": DocumentSourceType.CSV.value},
                    blocks=[
                        TableBlock(
                            headers=parsed["headers"],
                            rows=parsed["rows"],
                            caption=f"Extracted from {resolved_path.name}",
                            metadata={"source_path": str(resolved_path)},
                        ),
                    ],
                ),
            ],
            source=_build_source_record(
                resolved_path,
                DocumentSourceType.CSV,
                metadata={"row_count": parsed["row_count"], "column_count": parsed["column_count"], "ocr_used": False},
            ),
        )

    def _ingest_text(
        self,
        source_path: Path,
        workspace_path: Path,
        *,
        allowed_file_roots: tuple[Path | str, ...],
        base_dir: Path | str | None,
    ) -> _ParsedSource:
        resolved_path, _ = read_scoped_text_file(
            source_path,
            allowed_roots=allowed_file_roots,
            purpose="TXT document",
            media_type="text/plain",
            base_dir=base_dir,
        )
        parser_path = self._copy_to_workspace(resolved_path, workspace_path.with_suffix(".txt"))
        parsed = self._run_parser("TXT ingest parser", _parse_text_file, str(parser_path))
        section = DocumentSection(
            heading=resolved_path.stem,
            metadata={"source_path": str(resolved_path), "source_type": DocumentSourceType.TEXT.value},
            blocks=[
                ParagraphBlock(
                    text=paragraph,
                    metadata={"source_path": str(resolved_path)},
                )
                for paragraph in parsed["paragraphs"]
            ],
        )
        return _ParsedSource(
            sections=[section],
            source=_build_source_record(
                resolved_path,
                DocumentSourceType.TEXT,
                metadata={"paragraph_count": len(parsed["paragraphs"]), "ocr_used": False},
            ),
        )

    def _ingest_markdown(
        self,
        source_path: Path,
        workspace_path: Path,
        *,
        allowed_file_roots: tuple[Path | str, ...],
        base_dir: Path | str | None,
    ) -> _ParsedSource:
        resolved_path, _ = read_scoped_text_file(
            source_path,
            allowed_roots=allowed_file_roots,
            purpose="Markdown document",
            media_type="text/markdown",
            base_dir=base_dir,
        )
        parser_path = self._copy_to_workspace(resolved_path, workspace_path.with_suffix(resolved_path.suffix))
        parsed = self._run_parser("Markdown ingest parser", _parse_markdown_file, str(parser_path))
        sections = [
            DocumentSection(
                heading=section["heading"],
                level=section["level"],
                metadata={"source_path": str(resolved_path), "source_type": DocumentSourceType.MARKDOWN.value},
                blocks=[
                    ParagraphBlock(
                        text=block["text"],
                        style_tokens=block.get("style_tokens", []),
                        metadata={"source_path": str(resolved_path)},
                    )
                    if block["type"] == "paragraph"
                    else ListBlock(
                        ordered=block["ordered"],
                        items=block["items"],
                        metadata={"source_path": str(resolved_path)},
                    )
                    for block in section["blocks"]
                ],
            )
            for section in parsed["sections"]
        ]
        return _ParsedSource(
            sections=sections,
            source=_build_source_record(
                resolved_path,
                DocumentSourceType.MARKDOWN,
                metadata={
                    "section_count": parsed["section_count"],
                    "paragraph_count": parsed["paragraph_count"],
                    "list_count": parsed["list_count"],
                    "ocr_used": False,
                },
            ),
        )

    def _ingest_xlsx(
        self,
        source_path: Path,
        workspace_path: Path,
        *,
        allowed_file_roots: tuple[Path | str, ...],
        base_dir: Path | str | None,
    ) -> _ParsedSource:
        resolved_path = validate_scoped_binary_file(
            source_path,
            allowed_roots=allowed_file_roots,
            purpose="XLSX document",
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            base_dir=base_dir,
        )
        parser_path = self._copy_to_workspace(resolved_path, workspace_path.with_suffix(".xlsx"))
        parsed = self._run_parser("XLSX ingest parser", _parse_xlsx_file, str(parser_path))
        sections: list[DocumentSection] = []
        for sheet in parsed["sheets"]:
            sections.append(
                DocumentSection(
                    heading=f"{resolved_path.stem} - {sheet['title']}",
                    metadata={
                        "source_path": str(resolved_path),
                        "source_type": DocumentSourceType.XLSX.value,
                        "sheet_title": sheet["title"],
                    },
                    blocks=[
                        TableBlock(
                            headers=sheet["headers"],
                            rows=sheet["rows"],
                            caption=f"Sheet {sheet['title']}",
                            metadata={"sheet_title": sheet["title"], "source_path": str(resolved_path)},
                        ),
                    ],
                ),
            )
        return _ParsedSource(
            sections=sections,
            source=_build_source_record(
                resolved_path,
                DocumentSourceType.XLSX,
                metadata={"sheet_count": parsed["sheet_count"], "ocr_used": False},
            ),
        )

    def _ingest_docx(
        self,
        source_path: Path,
        workspace_path: Path,
        *,
        allowed_file_roots: tuple[Path | str, ...],
        base_dir: Path | str | None,
        ocr_status: OcrBackendStatus,
    ) -> _ParsedSource:
        resolved_path = validate_scoped_binary_file(
            source_path,
            allowed_roots=allowed_file_roots,
            purpose="DOCX document",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            base_dir=base_dir,
        )
        parser_path = self._copy_to_workspace(resolved_path, workspace_path.with_suffix(".docx"))
        parsed = self._run_parser("DOCX ingest parser", _parse_docx_file, str(parser_path))
        base_metadata = {"source_path": str(resolved_path), "source_type": DocumentSourceType.DOCX.value}

        sections: list[DocumentSection] = []
        current_heading = resolved_path.stem
        current_level = 1
        current_blocks: list[Any] = []

        def flush_current_section() -> None:
            if current_blocks:
                sections.append(
                    DocumentSection(
                        heading=current_heading,
                        level=current_level,
                        metadata={**base_metadata},
                        blocks=list(current_blocks),
                    ),
                )
                current_blocks.clear()

        for paragraph in parsed["paragraphs"]:
            text = paragraph["text"]
            style = paragraph.get("style")
            heading_level = _heading_level_from_style(style)
            if heading_level is not None:
                flush_current_section()
                current_heading = text
                current_level = heading_level
                continue
            current_blocks.append(
                ParagraphBlock(
                    text=text,
                    style_tokens=[style] if style else [],
                    metadata={"source_style": style} if style else {},
                ),
            )

        for table_index, table in enumerate(parsed["tables"], start=1):
            current_blocks.append(
                TableBlock(
                    headers=table["headers"],
                    rows=table["rows"],
                    caption=f"Table {table_index} from {resolved_path.name}",
                    metadata={"source_path": str(resolved_path), "table_index": table_index},
                ),
            )

        extracted_images = self._extract_docx_images(parser_path)
        ocr_used = False
        for image_index, image in enumerate(extracted_images, start=1):
            current_blocks.append(
                ImageBlock(
                    alt_text=image["file_name"],
                    caption=f"Embedded image {image_index} from {resolved_path.name}",
                    mime_type=image["mime_type"],
                    width=image["width"],
                    height=image["height"],
                    metadata={"source_name": image["source_name"], "source_path": str(resolved_path)},
                ),
            )
            ocr_result = perform_ocr_on_image_bytes(image["data"])
            if ocr_result.text:
                current_blocks.append(
                    ParagraphBlock(
                        text=ocr_result.text,
                        style_tokens=["ocr"],
                        metadata={"ocr_backend": ocr_result.backend_name, "source_name": image["source_name"]},
                    ),
                )
                ocr_used = True

        if parsed["inline_image_count"] and not extracted_images:
            current_blocks.append(
                CalloutBlock(
                    kind="note",
                    title="Inline Images",
                    body=f"Detected {parsed['inline_image_count']} inline image(s), but no extractable media payloads were recovered from the DOCX package.",
                ),
            )
        elif extracted_images and not ocr_used and not ocr_status.available:
            current_blocks.append(
                CalloutBlock(
                    kind="warning",
                    title="OCR",
                    body=f"Embedded images were extracted, but OCR is unavailable: {ocr_status.reason}",
                ),
            )

        if not current_blocks:
            current_blocks.append(CalloutBlock(kind="warning", title="Extraction", body="No extractable text, tables, or images were found."))
        flush_current_section()

        return _ParsedSource(
            sections=sections,
            source=_build_source_record(
                resolved_path,
                DocumentSourceType.DOCX,
                metadata={
                    "paragraph_count": len(parsed["paragraphs"]),
                    "table_count": len(parsed["tables"]),
                    "inline_image_count": parsed["inline_image_count"],
                    "extracted_image_count": len(extracted_images),
                    "ocr_used": ocr_used,
                },
            ),
        )

    def _ingest_pdf(
        self,
        source_path: Path,
        workspace_path: Path,
        *,
        allowed_file_roots: tuple[Path | str, ...],
        base_dir: Path | str | None,
        ocr_status: OcrBackendStatus,
    ) -> _ParsedSource:
        resolved_path = validate_scoped_binary_file(
            source_path,
            allowed_roots=allowed_file_roots,
            purpose="PDF document",
            media_type="application/pdf",
            base_dir=base_dir,
        )
        parser_path = self._copy_to_workspace(resolved_path, workspace_path.with_suffix(".pdf"))
        parsed = self._run_parser("PDF ingest parser", _parse_pdf_file, str(parser_path))
        sections: list[DocumentSection] = []
        ocr_page_count = 0
        for index, page in enumerate(parsed["pages"], start=1):
            blocks: list[Any] = []
            if page["paragraphs"]:
                blocks.extend(ParagraphBlock(text=segment) for segment in page["paragraphs"])

            if page["images"]:
                for image in page["images"]:
                    metadata = _read_image_metadata_bytes(image["data"], image["name"])
                    blocks.append(
                        ImageBlock(
                            alt_text=metadata["file_name"],
                            caption=f"Embedded PDF image on page {index}",
                            mime_type=metadata["mime_type"],
                            width=metadata["width"],
                            height=metadata["height"],
                            metadata={"source_path": str(resolved_path), "page_number": index},
                        ),
                    )

            page_ocr_texts: list[str] = []
            if not page["paragraphs"] and page["images"]:
                for image in page["images"]:
                    ocr_result = perform_ocr_on_image_bytes(image["data"])
                    if ocr_result.text:
                        page_ocr_texts.append(ocr_result.text)
                if page_ocr_texts:
                    ocr_page_count += 1
                    blocks = [
                        ParagraphBlock(
                            text=text,
                            style_tokens=["ocr"],
                            metadata={"ocr_backend": ocr_status.backend_name, "page_number": index},
                        )
                        for text in page_ocr_texts
                    ] + blocks
                else:
                    warning = page.get("warning", "No extractable text found on this page.")
                    if not ocr_status.available:
                        warning = f"{warning} OCR fallback unavailable: {ocr_status.reason}"
                    else:
                        warning = f"{warning} OCR did not detect text in the embedded page images."
                    blocks = [CalloutBlock(kind="warning", title="Extraction", body=warning)] + blocks

            if not blocks:
                warning = page.get("warning", "No extractable text found on this page.")
                if page["images"] and not ocr_status.available:
                    warning = f"{warning} OCR fallback unavailable: {ocr_status.reason}"
                elif page["images"] and not page_ocr_texts:
                    warning = f"{warning} OCR did not detect text in the embedded page images."
                blocks = [CalloutBlock(kind="warning", title="Extraction", body=warning)]

            sections.append(
                DocumentSection(
                    heading=f"{resolved_path.stem} - Page {index}",
                    metadata={
                        "source_path": str(resolved_path),
                        "source_type": DocumentSourceType.PDF.value,
                        "page_number": index,
                        "ocr_used": bool(page_ocr_texts),
                    },
                    blocks=blocks,
                ),
            )
        return _ParsedSource(
            sections=sections,
            source=_build_source_record(
                resolved_path,
                DocumentSourceType.PDF,
                metadata={
                    "page_count": parsed["page_count"],
                    "text_page_count": parsed["text_page_count"],
                    "image_count": parsed["image_count"],
                    "ocr_used": ocr_page_count > 0,
                    "ocr_page_count": ocr_page_count,
                },
            ),
        )

    def _ingest_image(
        self,
        source_path: Path,
        workspace_path: Path,
        *,
        allowed_file_roots: tuple[Path | str, ...],
        base_dir: Path | str | None,
        ocr_status: OcrBackendStatus,
    ) -> _ParsedSource:
        resolved_path = validate_scoped_image_file(
            source_path,
            allowed_roots=allowed_file_roots,
            purpose="Image document",
            base_dir=base_dir,
        )
        parser_path = self._copy_to_workspace(resolved_path, workspace_path.with_suffix(resolved_path.suffix))
        image_metadata = self._extract_image_metadata(parser_path, resolved_path)
        raw = parser_path.read_bytes()
        ocr_result = perform_ocr_on_image_bytes(raw)
        image_metadata["ocr_text"] = ocr_result.text
        image_metadata["ocr_backend"] = ocr_result.backend_name
        image_metadata["ocr_reason"] = ocr_result.reason
        return _ParsedSource(
            sections=[
                self._image_section(
                    resolved_path.stem,
                    [image_metadata],
                    source_type=DocumentSourceType.IMAGE,
                    ocr_status=ocr_status,
                ),
            ],
            source=_build_source_record(
                resolved_path,
                DocumentSourceType.IMAGE,
                metadata={
                    **image_metadata,
                    "ocr_used": bool(ocr_result.text),
                    "ocr_text_characters": len(ocr_result.text or ""),
                },
            ),
        )

    def _ingest_image_bundle(
        self,
        source_path: Path,
        workspace_path: Path,
        *,
        allowed_file_roots: tuple[Path | str, ...],
        base_dir: Path | str | None,
        ocr_status: OcrBackendStatus,
    ) -> _ParsedSource:
        image_paths = []
        for child in sorted(source_path.rglob("*")):
            if not child.is_file():
                continue
            try:
                image_paths.append(
                    validate_scoped_image_file(
                        child,
                        allowed_roots=allowed_file_roots,
                        purpose="Image bundle item",
                        base_dir=base_dir,
                    ),
                )
            except UnsupportedMediaTypeError:
                continue
        if not image_paths:
            raise DocumentValidationError("Image bundle did not contain supported image files.", details={"path": str(source_path)})
        workspace_path.mkdir(parents=True, exist_ok=True)
        image_metadata = []
        ocr_used_count = 0
        for image_path in image_paths:
            copied = self._copy_to_workspace(image_path, workspace_path / image_path.name)
            metadata = self._extract_image_metadata(copied, image_path)
            ocr_result = perform_ocr_on_image_bytes(copied.read_bytes())
            metadata["ocr_text"] = ocr_result.text
            metadata["ocr_backend"] = ocr_result.backend_name
            metadata["ocr_reason"] = ocr_result.reason
            image_metadata.append(metadata)
            if ocr_result.text:
                ocr_used_count += 1
        return _ParsedSource(
            sections=[
                self._image_section(
                    source_path.name,
                    image_metadata,
                    source_type=DocumentSourceType.IMAGE_BUNDLE,
                    ocr_status=ocr_status,
                ),
            ],
            source=_build_source_record(
                source_path,
                DocumentSourceType.IMAGE_BUNDLE,
                metadata={"image_count": len(image_paths), "ocr_used": ocr_used_count > 0, "ocr_image_count": ocr_used_count},
            ),
        )

    def _image_section(
        self,
        heading: str,
        image_metadata: list[dict[str, Any]],
        *,
        source_type: DocumentSourceType,
        ocr_status: OcrBackendStatus,
    ) -> DocumentSection:
        rows = [
            [
                image["file_name"],
                image["format"],
                str(image["width"]),
                str(image["height"]),
                "yes" if image.get("ocr_text") else "no",
            ]
            for image in image_metadata
        ]
        ocr_image_count = sum(1 for image in image_metadata if image.get("ocr_text"))
        if ocr_image_count:
            ocr_callout = CalloutBlock(
                kind="success",
                title="OCR",
                body=f"OCR extracted text from {ocr_image_count} of {len(image_metadata)} image(s).",
            )
        elif ocr_status.available:
            ocr_callout = CalloutBlock(
                kind="note",
                title="OCR",
                body="OCR backend is available, but no text was detected in the supplied images.",
            )
        else:
            ocr_callout = CalloutBlock(
                kind="warning",
                title="OCR",
                body=f"OCR is unavailable; returning image metadata and file references only. {ocr_status.reason}",
            )

        blocks: list[Any] = [
            ocr_callout,
            TableBlock(
                headers=["File", "Format", "Width", "Height", "OCR"],
                rows=rows,
                caption="Extracted image metadata",
            ),
        ]
        for image in image_metadata:
            blocks.append(
                ImageBlock(
                    alt_text=image["file_name"],
                    path=image.get("path"),
                    caption=f"{image['format']} {image['width']}x{image['height']}",
                    mime_type=image.get("mime_type"),
                    width=image["width"],
                    height=image["height"],
                    metadata={"source_path": image.get("path")},
                ),
            )
            if image.get("ocr_text"):
                blocks.append(
                    ParagraphBlock(
                        text=image["ocr_text"],
                        style_tokens=["ocr"],
                        metadata={"ocr_backend": image.get("ocr_backend"), "file_name": image["file_name"]},
                    ),
                )
        return DocumentSection(
            heading=heading,
            metadata={"source_type": source_type.value},
            blocks=blocks,
        )

    def _extract_image_metadata(self, parser_path: Path, original_path: Path) -> dict[str, Any]:
        image_metadata = self._run_parser("Image metadata parser", _read_image_metadata, str(parser_path))
        return {
            "path": str(original_path),
            "file_name": original_path.name,
            "format": image_metadata["format"] or original_path.suffix.lstrip(".").upper() or "UNKNOWN",
            "width": image_metadata["width"],
            "height": image_metadata["height"],
            "mime_type": _guess_image_media_type(original_path.name, image_metadata["format"]),
        }

    def _extract_docx_images(self, parser_path: Path) -> list[dict[str, Any]]:
        images: list[dict[str, Any]] = []
        with zipfile.ZipFile(parser_path) as archive:
            for name in sorted(member for member in archive.namelist() if member.startswith("word/media/") and not member.endswith("/")):
                raw = archive.read(name)
                metadata = _read_image_metadata_bytes(raw, Path(name).name)
                images.append({**metadata, "data": raw, "source_name": name})
        return images

    def _copy_to_workspace(self, source_path: Path, target_path: Path) -> Path:
        target_path.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(source_path, target_path)
        return target_path

    def _run_parser(self, operation: str, parser, parser_path: str) -> dict[str, Any]:
        return run_in_subprocess(
            parser,
            parser_path,
            operation=operation,
            timeout_seconds=self.sandbox_timeout_seconds,
            enabled=self.sandbox_enabled,
            clear_environment=self.sandbox_clear_environment,
            workspace_root=self.workspace_root / "parser-sandbox",
        )

    def _build_chunks(self, document: DocumentIR) -> list[DocumentChunk]:
        chunks: list[DocumentChunk] = []
        for section_index, section in enumerate(document.sections):
            section_metadata = section.metadata or {}
            source_path = section_metadata.get("source_path")
            source_type = _coerce_source_type(section_metadata.get("source_type"))
            source_id = str(section_metadata.get("source_id") or _stable_source_id(str(source_path or f"section-{section_index}")))
            section_id = str(section_metadata.get("section_id") or f"{source_id}-sec-{section_index + 1:04d}")
            source_name = (
                str(section_metadata.get("source_name"))
                if section_metadata.get("source_name") is not None
                else (Path(str(source_path)).name if source_path else None)
            )
            source_label = (
                str(section_metadata.get("source_label"))
                if section_metadata.get("source_label") is not None
                else (source_name or f"Source {section_index + 1}")
            )
            section_label = (
                str(section_metadata.get("section_label"))
                if section_metadata.get("section_label") is not None
                else _build_section_label(section, source_label=source_label, source_section_index=section_index)
            )
            parts = [part for part in _section_to_chunk_parts(section) if part]
            if not parts:
                continue
            current_parts: list[str] = []
            current_length = 0
            section_chunk_index = 0
            for part in parts:
                for segment in _split_text_segment(part, max_characters=self.chunk_max_characters):
                    projected = current_length + len(segment) + (2 if current_parts else 0)
                    if current_parts and projected > self.chunk_max_characters:
                        text = "\n\n".join(current_parts).strip()
                        document_chunk_index = len(chunks)
                        section_chunk_index += 1
                        chunks.append(
                            DocumentChunk(
                                chunk_id=f"{section_id}-chunk-{section_chunk_index:04d}",
                                text=text,
                                source_id=source_id,
                                section_id=section_id,
                                source_label=source_label,
                                section_label=section_label,
                                section_heading=section.heading,
                                section_level=section.level,
                                source_name=source_name,
                                source_path=source_path,
                                source_type=source_type,
                                metadata={
                                    "chunk_index": document_chunk_index,
                                    "section_chunk_index": section_chunk_index - 1,
                                    "section_index": section_metadata.get("section_index", section_index),
                                    "source_index": section_metadata.get("source_index"),
                                    "source_section_index": section_metadata.get("source_section_index"),
                                    "section_level": section.level,
                                    "char_count": len(text),
                                    **_extract_chunk_provenance(section_metadata),
                                },
                            ),
                        )
                        current_parts = []
                        current_length = 0
                    current_parts.append(segment)
                    current_length += len(segment) + (2 if current_parts[:-1] else 0)
            if current_parts:
                text = "\n\n".join(current_parts).strip()
                document_chunk_index = len(chunks)
                section_chunk_index += 1
                chunks.append(
                    DocumentChunk(
                        chunk_id=f"{section_id}-chunk-{section_chunk_index:04d}",
                        text=text,
                        source_id=source_id,
                        section_id=section_id,
                        source_label=source_label,
                        section_label=section_label,
                        section_heading=section.heading,
                        section_level=section.level,
                        source_name=source_name,
                        source_path=source_path,
                        source_type=source_type,
                        metadata={
                            "chunk_index": document_chunk_index,
                            "section_chunk_index": section_chunk_index - 1,
                            "section_index": section_metadata.get("section_index", section_index),
                            "source_index": section_metadata.get("source_index"),
                            "source_section_index": section_metadata.get("source_section_index"),
                            "section_level": section.level,
                            "char_count": len(text),
                            **_extract_chunk_provenance(section_metadata),
                        },
                    ),
                )
        return chunks

    def _default_title(self, paths: list[Path]) -> str:
        if len(paths) == 1:
            return paths[0].stem if paths[0].is_file() else paths[0].name
        return "Ingested Document Bundle"

    def _publish(self, event_type: EventType, payload: dict[str, object]) -> None:
        if self.event_bus is None:
            return
        self.event_bus.publish_threadsafe(
            StreamEvent(type=event_type, scope=EventScope.REQUEST, payload=payload),
        )


def _heading_level_from_style(style_name: str | None) -> int | None:
    if not style_name:
        return None
    normalized = style_name.casefold().strip()
    if normalized == "title":
        return 1
    match = re.search(r"heading\s*([1-9])", normalized)
    if match is None:
        return None
    return int(match.group(1))


def _join_wrapped_lines(lines: list[str]) -> str:
    return " ".join(line.strip() for line in lines if line.strip()).strip()


def _paragraphs_from_text(text: str) -> list[str]:
    return [
        _join_wrapped_lines(block.splitlines())
        for block in re.split(r"\n\s*\n", text)
        if _join_wrapped_lines(block.splitlines())
    ]


def _guess_image_media_type(file_name: str, image_format: str | None = None) -> str:
    suffix = Path(file_name).suffix.casefold()
    if suffix == ".png" or (image_format or "").casefold() == "png":
        return "image/png"
    if suffix in {".jpg", ".jpeg"} or (image_format or "").casefold() == "jpeg":
        return "image/jpeg"
    if suffix == ".gif" or (image_format or "").casefold() == "gif":
        return "image/gif"
    return "image/*"


def _stable_source_id(path: str) -> str:
    digest = hashlib.sha1(path.encode("utf-8")).hexdigest()[:12]
    return f"src-{digest}"


def _guess_source_media_type(path: str, source_type: DocumentSourceType) -> str | None:
    if source_type == DocumentSourceType.TEXT:
        return "text/plain"
    if source_type == DocumentSourceType.MARKDOWN:
        return "text/markdown"
    if source_type == DocumentSourceType.CSV:
        return "text/csv"
    if source_type == DocumentSourceType.XLSX:
        return "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    if source_type == DocumentSourceType.DOCX:
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if source_type == DocumentSourceType.PDF:
        return "application/pdf"
    if source_type == DocumentSourceType.IMAGE:
        return mimetypes.guess_type(path)[0] or _guess_image_media_type(path)
    if source_type == DocumentSourceType.IMAGE_BUNDLE:
        return None
    return mimetypes.guess_type(path)[0]


def _standardize_source_record(source: IngestedDocumentSource, *, source_index: int) -> IngestedDocumentSource:
    return source.model_copy(
        update={
            "metadata": {
                **source.metadata,
                "source_index": source_index,
            },
        },
    )


def _build_source_record(
    path: Path | str,
    source_type: DocumentSourceType,
    *,
    metadata: dict[str, Any],
) -> IngestedDocumentSource:
    resolved_path = str(Path(path).resolve(strict=False))
    source_name = Path(resolved_path).name or resolved_path
    return IngestedDocumentSource(
        source_id=_stable_source_id(resolved_path),
        path=resolved_path,
        source_type=source_type,
        source_name=source_name,
        source_label=source_name,
        media_type=_guess_source_media_type(resolved_path, source_type),
        metadata=metadata,
    )


def _standardize_sections(
    sections: list[DocumentSection],
    *,
    source: IngestedDocumentSource,
    source_index: int,
    global_section_offset: int,
) -> list[DocumentSection]:
    updated_sections: list[DocumentSection] = []
    for source_section_index, section in enumerate(sections):
        section_id = f"{source.source_id}-sec-{source_section_index + 1:04d}"
        section_label = _build_section_label(section, source_label=source.source_label, source_section_index=source_section_index)
        updated_sections.append(
            section.model_copy(
                update={
                    "metadata": {
                        **section.metadata,
                        "source_id": source.source_id,
                        "source_index": source_index,
                        "source_name": source.source_name,
                        "source_label": source.source_label,
                        "source_path": source.path,
                        "source_type": source.source_type.value,
                        "source_media_type": source.media_type,
                        "section_id": section_id,
                        "section_index": global_section_offset + source_section_index,
                        "source_section_index": source_section_index,
                        "section_label": section_label,
                        "section_heading": section.heading,
                        "section_level": section.level,
                    },
                },
            ),
        )
    return updated_sections


def _build_section_label(section: DocumentSection, *, source_label: str, source_section_index: int) -> str:
    metadata = section.metadata or {}
    page_number = metadata.get("page_number")
    if page_number is not None:
        return f"{source_label} / Page {page_number}"
    sheet_title = metadata.get("sheet_title")
    if isinstance(sheet_title, str) and sheet_title.strip():
        return f"{source_label} / Sheet {sheet_title.strip()}"
    heading = (section.heading or "").strip()
    if heading:
        return f"{source_label} / {heading}"
    return f"{source_label} / Section {source_section_index + 1}"


def _section_to_chunk_parts(section: DocumentSection) -> list[str]:
    parts: list[str] = []
    if section.heading:
        parts.append(f"Section: {section.heading}")
    for block in section.blocks:
        text = _block_to_text(block)
        if text:
            parts.append(text)
    return parts


def _block_to_text(block: Any) -> str:
    if isinstance(block, ParagraphBlock):
        return block.text.strip()
    if isinstance(block, ListBlock):
        prefix = "1. " if block.ordered else "- "
        return "\n".join(f"{prefix}{item}" for item in block.items if item.strip())
    if isinstance(block, TableBlock):
        lines: list[str] = []
        if block.caption:
            lines.append(block.caption)
        if block.headers:
            lines.append(" | ".join(block.headers))
        lines.extend(" | ".join(row) for row in block.rows)
        return "\n".join(lines).strip()
    if isinstance(block, CalloutBlock):
        title = f"{block.kind.upper()}: {block.title}" if block.title else block.kind.upper()
        return f"{title}\n{block.body}".strip()
    if isinstance(block, ImageBlock):
        caption = f" ({block.caption})" if block.caption else ""
        return f"Image: {block.alt_text}{caption}"
    return ""


def _split_text_segment(text: str, *, max_characters: int) -> list[str]:
    normalized = text.strip()
    if not normalized:
        return []
    if len(normalized) <= max_characters:
        return [normalized]
    parts: list[str] = []
    current = ""
    for sentence in re.split(r"(?<=[.!?])\s+", normalized):
        sentence = sentence.strip()
        if not sentence:
            continue
        projected = f"{current} {sentence}".strip()
        if current and len(projected) > max_characters:
            parts.append(current)
            current = sentence
            continue
        if len(sentence) > max_characters:
            words = sentence.split()
            current_word_group = ""
            for word in words:
                projected_word_group = f"{current_word_group} {word}".strip()
                if current_word_group and len(projected_word_group) > max_characters:
                    parts.append(current_word_group)
                    current_word_group = word
                else:
                    current_word_group = projected_word_group
            if current_word_group:
                if current:
                    parts.append(current)
                    current = current_word_group
                else:
                    current = current_word_group
            continue
        current = projected
    if current:
        parts.append(current)
    return parts


def _extract_chunk_provenance(section_metadata: dict[str, Any]) -> dict[str, Any]:
    return {
        key: section_metadata[key]
        for key in (
            "page_number",
            "sheet_title",
            "ocr_used",
            "ocr_backend",
            "source_media_type",
        )
        if key in section_metadata and section_metadata[key] is not None
    }


def _coerce_source_type(value: Any) -> DocumentSourceType | None:
    if value is None:
        return None
    try:
        return DocumentSourceType(value)
    except ValueError:
        return None
